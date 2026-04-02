"""
Export patent application to Word (.docx), PDF, and Markdown formats.
"""

from __future__ import annotations
import io
from typing import TYPE_CHECKING

from .models import ApplicationResult

if TYPE_CHECKING:
    pass


def export_markdown(result: ApplicationResult) -> str:
    """Export the application as a Markdown document."""
    lines = []
    lines.append(f"# {result.title}")
    lines.append("")

    if result.cross_references:
        lines.append("## Cross-Reference to Related Applications")
        lines.append("")
        lines.append(result.cross_references)
        lines.append("")

    lines.append("## Background of the Invention")
    lines.append("")
    lines.append(result.background)
    lines.append("")

    lines.append("## Summary of the Invention")
    lines.append("")
    lines.append(result.summary)
    lines.append("")

    if result.figure_descriptions:
        lines.append("## Brief Description of the Drawings")
        lines.append("")
        lines.append(result.figure_descriptions)
        lines.append("")

    lines.append("## Detailed Description of the Preferred Embodiments")
    lines.append("")
    lines.append(result.detailed_description)
    lines.append("")

    lines.append("## Claims")
    lines.append("")
    lines.append(result.claims_text)
    lines.append("")

    lines.append("## Abstract of the Disclosure")
    lines.append("")
    lines.append(result.abstract)
    lines.append("")

    return "\n".join(lines)


def export_docx(result: ApplicationResult) -> bytes:
    """Export the application as a Word (.docx) document."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title_para = doc.add_heading(result.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    section_data = [
        ("Cross-Reference to Related Applications", result.cross_references),
        ("Background of the Invention", result.background),
        ("Summary of the Invention", result.summary),
        ("Brief Description of the Drawings", result.figure_descriptions),
        ("Detailed Description of the Preferred Embodiments", result.detailed_description),
    ]

    for heading, content in section_data:
        if not content.strip():
            continue
        doc.add_heading(heading, level=1)
        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

    # Claims
    doc.add_heading("Claims", level=1)
    for para_text in result.claims_text.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    # Abstract
    doc.add_heading("Abstract of the Disclosure", level=1)
    doc.add_paragraph(result.abstract)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(result: ApplicationResult) -> bytes:
    """Export the application as a PDF document via WeasyPrint."""
    from weasyprint import HTML

    md = export_markdown(result)

    # Convert markdown-ish text to simple HTML
    html_sections = []
    for line in md.split("\n"):
        if line.startswith("# "):
            html_sections.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("## "):
            html_sections.append(f"<h2>{line[3:]}</h2>")
        elif line.strip() == "":
            html_sections.append("")
        else:
            html_sections.append(f"<p>{line}</p>")

    html_content = f"""<!DOCTYPE html>
<html>
<head>
<style>
    body {{
        font-family: 'Times New Roman', Times, serif;
        font-size: 12pt;
        line-height: 1.5;
        margin: 1in;
        color: #000;
    }}
    h1 {{
        text-align: center;
        font-size: 14pt;
        margin-top: 24pt;
    }}
    h2 {{
        font-size: 12pt;
        font-weight: bold;
        text-transform: uppercase;
        margin-top: 18pt;
    }}
    p {{
        text-align: justify;
        text-indent: 0.5in;
        margin: 6pt 0;
    }}
</style>
</head>
<body>
{"".join(html_sections)}
</body>
</html>"""

    pdf_bytes = HTML(string=html_content).write_pdf()
    return pdf_bytes
